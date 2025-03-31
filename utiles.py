# utils.py
import re
from typing import Tuple, List
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from io import BytesIO

def clean_indices(spans: List[Tuple[int, int]]) -> List[Tuple[int, int]]:
    """Фильтрует перекрывающиеся индексы."""
    sorted_indices = sorted(spans, key=lambda x: x[0])
    filtered_indices = []
    for span in sorted_indices:
        if not any(start <= span[0] and end >= span[1] for start, end in filtered_indices):
            filtered_indices.append(span)
    return filtered_indices

def insert_highlights(text: str, fix_spans: List[Tuple[int, int]], color: str = 'blue') -> str:
    """
    Вставляет HTML-теги подсветки для исправлений в тексте.
    
    :param text: Исходный текст.
    :param fix_spans: Список кортежей (start, end) для исправлений.
    :param color: Цвет подсветки (по умолчанию 'blue').
    :return: Текст с вставленными HTML-тегами подсветки.
    """
    fix_spans = clean_indices(fix_spans)
    fix_spans = sorted(fix_spans, key=lambda x: x[0], reverse=True)
    
    for start, end in fix_spans:
        highlight_tag = f'<mark style="background-color: {color};">{text[start:end]}</mark>'
        text = text[:start] + highlight_tag + text[end:]
    
    return text

from docx import Document
import re

def convert_docx_to_html(doc: Document) -> str:
    """
    Конвертирует docx в HTML с подсветкой исправлений.
    """
    html = ""

    for para in doc.paragraphs:
        para_html = ""
        for run in para.runs:
            text = run.text
            if run.font.highlight_color:
                # Определяем цвет подсветки
                if run.font.highlight_color == 1:  # WD_COLOR_INDEX.YELLOW
                    color = 'blue'
                else:
                    color = 'blue'  # Можно добавить другие цвета по необходимости
                para_html += f'<mark style="background-color: {color};">{text}</mark>'
            else:
                para_html += text
        html += f"<p>{para_html}</p>"

    # Обработка таблиц (опционально)
    for table in doc.tables:
        html += "<table>"
        for row in table.rows:
            html += "<tr>"
            for cell in row.cells:
                cell_html = ""
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text
                        if run.font.highlight_color:
                            if run.font.highlight_color == 1:  # WD_COLOR_INDEX.BLUE
                                color = 'blue'
                            else:
                                color = 'blue'
                            cell_html += f'<mark style="background-color: {color};">{text}</mark>'
                        else:
                            cell_html += text
                    cell_html += "<br>"  # Разделение параграфов в ячейке
                html += f"<td>{cell_html}</td>"
            html += "</tr>"
        html += "</table>"

    return html

def apply_highlights_to_docx(doc: Document, fix_spans: List[Tuple[int, int]], color: str = 'blue'):
    """
    Применяет подсветку к указанным спанам в документе DOCX.
    
    :param doc: Объект Document.
    :param fix_spans: Список спанов (start, end) для подсветки.
    :param color: Цвет подсветки (по умолчанию 'yellow').
    """
    fix_spans = clean_indices(fix_spans or [])
    current_pos = 0

    for para in doc.paragraphs:
        para_length = len(para.text)
        spans_in_para = []

        for start, end in fix_spans:
            if start >= current_pos and end <= current_pos + para_length:
                spans_in_para.append((start - current_pos, end - current_pos))
        
        for run in para.runs:
            run_length = len(run.text)
            run_start = current_pos
            run_end = current_pos + run_length

            for span_start, span_end in spans_in_para:
                if span_start < run_length and span_end > 0:
                    highlight_start = max(span_start, 0)
                    highlight_end = min(span_end, run_length)
                    if highlight_start < highlight_end:
                        before = run.text[:highlight_start]
                        highlight_text = run.text[highlight_start:highlight_end]
                        after = run.text[highlight_end:]

                        run.text = before

                        if highlight_text:
                            highlighted_run = para.add_run(highlight_text)
                            highlighted_run.font.highlight_color = WD_COLOR_INDEX.BLUE  # Выберите нужный цвет

                        if after:
                            after_run = para.add_run(after)
                        break  # Переходим к следующему run после изменения

        current_pos += para_length + 1  # +1 для перевода строки

def highlight_text(text: str, spans: List[Tuple[int, int]], color: str) -> str:
    """
    Вставляет HTML-теги подсветки с указанным цветом на заданных спанах.
    
    :param text: Исходный текст.
    :param spans: Список кортежей (start, end) для подсветки.
    :param color: Цвет подсветки (например, 'red' или 'blue').
    :return: Текст с вставленными HTML-тегами подсветки.
    """
    spans = clean_indices(spans)
    spans = sorted(spans, key=lambda x: x[0], reverse=True)
    
    for start, end in spans:
        highlight_tag = f'<mark style="background-color: {color};">{text[start:end]}</mark>'
        text = text[:start] + highlight_tag + text[end:]
    
    return text
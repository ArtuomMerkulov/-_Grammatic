# task.py

import re
from typing import Tuple, List
from docx import Document
from docx.enum.text import WD_COLOR_INDEX  # Для подсветки
import logging

logger = logging.getLogger(__name__)

def task3(doc: Document) -> Tuple[Document, List[str]]:
    """
    Удаляет курсивное начертание из текста.
    Возвращает обновленный документ и историю изменений.
    """
    history = ["3. Курсив в докладной записке не используется"]
    
    for para in doc.paragraphs:
        for run in para.runs:
            if run.italic:
                original_text = run.text
                run.italic = False
                run.font.highlight_color = WD_COLOR_INDEX.BLUE  # Установка подсветки
                logger.debug(f"task3 изменено: '{original_text}' -> '{run.text}'")
    
    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.italic:
                            original_text = run.text
                            run.italic = False
                            run.font.highlight_color = WD_COLOR_INDEX.BLUE  # Установка подсветки
                            logger.debug(f"task3 в таблице изменено: '{original_text}' -> '{run.text}'")
    
    logger.info(f"task3 выполнена: {history}")
    return doc, history

def task10(doc: Document) -> Tuple[Document, List[str]]:
    """
    Удаляет пробелы при указании номера документа или пункта.
    Возвращает обновленный документ и историю изменений.
    """
    history = ["10. При указании номера документа или его пункта пробел не ставится"]
    patterns = [
        (r"(п\.)\s*(\d[\d.]*)", r"\1\2"),  # п.1.2.3
        (r"(пп\.)\s*(\d[\d.]*)\s*-\s*(\d[\d.]*)", r"\1\2-\3"),  # пп.1.2.3-1.2.5
        (r"№\s*(\d+)", r"№\1"),  # № 3 -> №3
    ]

    for para in doc.paragraphs:
        for run in para.runs:
            original_text = run.text
            new_text = run.text
            changed = False
            for pattern, replace_with in patterns:
                if re.search(pattern, new_text):
                    new_text = re.sub(pattern, replace_with, new_text)
                    changed = True
            if changed and original_text != new_text:
                run.text = new_text
                run.font.highlight_color = WD_COLOR_INDEX.BLUE
                logger.debug(f"task10 изменено: '{original_text}' -> '{new_text}'")
    
    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        original_text = run.text
                        new_text = run.text
                        changed = False
                        for pattern, replace_with in patterns:
                            if re.search(pattern, new_text):
                                new_text = re.sub(pattern, replace_with, new_text)
                                changed = True
                        if changed and original_text != new_text:
                            run.text = new_text
                            run.font.highlight_color = WD_COLOR_INDEX.BLUE
                            logger.debug(f"task10 в таблице изменено: '{original_text}' -> '{new_text}'")
    
    logger.info(f"task10 выполнена: {history}")
    return doc, history

def task12(doc: Document) -> Tuple[Document, List[str]]:
    """
    Заменяет кавычки типа "..." на формат «...».
    Возвращает обновленный документ и историю изменений.
    """
    history = ["12. Формат кавычек - «...»"]
    pattern = r'"(.*?)"'
    replace_with = r'«\1»'

    for para in doc.paragraphs:
        for run in para.runs:
            if '"' in run.text:
                original_text = run.text
                new_text = re.sub(pattern, replace_with, run.text)
                if original_text != new_text:
                    run.text = new_text
                    run.font.highlight_color = WD_COLOR_INDEX.BLUE
                    logger.debug(f'task12 изменено: "{original_text}" -> "{new_text}"')
    
    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if '"' in run.text:
                            original_text = run.text
                            new_text = re.sub(pattern, replace_with, run.text)
                            if original_text != new_text:
                                run.text = new_text
                                run.font.highlight_color = WD_COLOR_INDEX.BLUE
                                logger.debug(f'task12 в таблице изменено: "{original_text}" -> "{new_text}"')
    
    logger.info(f"task12 выполнена: {history}")
    return doc, history

def task13(doc: Document) -> Tuple[Document, List[str]]:
    """
    Форматирует сокращения: убирает точки у 'млн' и 'млрд', добавляет точки у 'тыс.' и 'руб.'.
    Возвращает обновленный документ и историю изменений.
    """
    history = ["13. Форматирование сокращений: 'млн' и 'млрд' без точек, 'тыс.' и 'руб.' с точками"]
    patterns = [
        (r'\b(млн|млрд)\b\.?', r'\1'),  # Убираем точки
        (r'\bтыс\b', 'тыс.'),           # Добавляем точку
        (r'\bруб\b', 'руб.'),           # Добавляем точку
    ]

    for para in doc.paragraphs:
        for run in para.runs:
            original_text = run.text
            new_text = run.text
            changed = False
            for pattern, replace_with in patterns:
                if re.search(pattern, new_text):
                    new_text = re.sub(pattern, replace_with, new_text)
                    changed = True
            if changed and original_text != new_text:
                run.text = new_text
                run.font.highlight_color = WD_COLOR_INDEX.BLUE
                logger.debug(f"task13 изменено: '{original_text}' -> '{new_text}'")

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        original_text = run.text
                        new_text = run.text
                        changed = False
                        for pattern, replace_with in patterns:
                            if re.search(pattern, new_text):
                                new_text = re.sub(pattern, replace_with, new_text)
                                changed = True
                        if changed and original_text != new_text:
                            run.text = new_text
                            run.font.highlight_color = WD_COLOR_INDEX.BLUE
                            logger.debug(f"task13 в таблице изменено: '{original_text}' -> '{new_text}'")
    
    logger.info(f"task13 выполнена: {history}")
    return doc, history

def task14(doc: Document) -> Tuple[Document, List[str]]:
    """
    Удаляет пробелы вокруг знака '/'.
    Возвращает обновленный документ и историю изменений.
    """
    history = ["14. Пробелы вокруг знака '/' удалены"]
    pattern = r'\s*/\s*'
    replace_with = '/'

    for para in doc.paragraphs:
        for run in para.runs:
            if re.search(pattern, run.text):
                original_text = run.text
                new_text = re.sub(pattern, replace_with, run.text)
                if original_text != new_text:
                    run.text = new_text
                    run.font.highlight_color = WD_COLOR_INDEX.BLUE
                    logger.debug(f"task14 изменено: '{original_text}' -> '{new_text}'")

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if re.search(pattern, run.text):
                            original_text = run.text
                            new_text = re.sub(pattern, replace_with, run.text)
                            if original_text != new_text:
                                run.text = new_text
                                run.font.highlight_color = WD_COLOR_INDEX.BLUE
                                logger.debug(f"task14 в таблице изменено: '{original_text}' -> '{new_text}'")
    
    logger.info(f"task14 выполнена: {history}")
    return doc, history

def task15(doc: Document) -> Tuple[Document, List[str]]:
    """
    Удаляет пробел перед знаком '%'.
    Возвращает обновленный документ и историю изменений.
    """
    history = ["15. Удаление пробела перед знаком '%'"]
    pattern = r'\s*%\s*'
    replace_with = '%'

    for para in doc.paragraphs:
        for run in para.runs:
            if re.search(pattern, run.text):
                original_text = run.text
                # Чтобы избежать потери контекста, сохраняем пробелы только перед %
                new_text = re.sub(r'\s+%', '%', run.text)
                new_text = re.sub(r'%\s+', '%', new_text)
                if original_text != new_text:
                    run.text = new_text
                    run.font.highlight_color = WD_COLOR_INDEX.BLUE
                    logger.debug(f"task15 изменено: '{original_text}' -> '{new_text}'")

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if re.search(pattern, run.text):
                            original_text = run.text
                            new_text = re.sub(r'\s+%', '%', run.text)
                            new_text = re.sub(r'%\s+', '%', new_text)
                            if original_text != new_text:
                                run.text = new_text
                                run.font.highlight_color = WD_COLOR_INDEX.BLUE
                                logger.debug(f"task15 в таблице изменено: '{original_text}' -> '{new_text}'")
    
    logger.info(f"task15 выполнена: {history}")
    return doc, history

def task19(doc: Document) -> Tuple[Document, List[str]]:
    """
    Правильное использование тире и дефисов.
    Возвращает обновленный документ и историю изменений.
    """
    history = ["19. Правильное использование тире и дефисов"]
    pattern_dash = r'\s*—\s*'  # Тире (длинное)
    replace_dash = ' — '       # Тире с пробелами
    pattern_hyphen = r'(?<!\s)-(?!=\s)'  # Дефис без пробелов
    replace_hyphen = '-'        # Дефис без пробелов

    for para in doc.paragraphs:
        for run in para.runs:
            original_text = run.text
            changed = False

            # Обработка тире
            if re.search(pattern_dash, run.text):
                new_text = re.sub(pattern_dash, replace_dash, run.text)
                if new_text != run.text:
                    run.text = new_text
                    run.font.highlight_color = WD_COLOR_INDEX.BLUE
                    changed = True
                    logger.debug(f"task19 (тире) изменено: '{original_text}' -> '{new_text}'")

            original_text = run.text  # Обновляем оригинальный текст после тире
            # Обработка дефиса
            if re.search(pattern_hyphen, run.text):
                new_text = re.sub(pattern_hyphen, replace_hyphen, run.text)
                if new_text != run.text:
                    run.text = new_text
                    run.font.highlight_color = WD_COLOR_INDEX.BLUE
                    changed = True
                    logger.debug(f"task19 (дефис) изменено: '{original_text}' -> '{new_text}'")
    
    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        original_text = run.text
                        changed = False

                        # Обработка тире
                        if re.search(pattern_dash, run.text):
                            new_text = re.sub(pattern_dash, replace_dash, run.text)
                            if new_text != run.text:
                                run.text = new_text
                                run.font.highlight_color = WD_COLOR_INDEX.BLUE
                                changed = True
                                logger.debug(f"task19 в таблице (тире) изменено: '{original_text}' -> '{new_text}'")

                        original_text = run.text  # Обновляем оригинальный текст после тире
                        # Обработка дефиса
                        if re.search(pattern_hyphen, run.text):
                            new_text = re.sub(pattern_hyphen, replace_hyphen, run.text)
                            if new_text != run.text:
                                run.text = new_text
                                run.font.highlight_color = WD_COLOR_INDEX.BLUE
                                changed = True
                                logger.debug(f"task19 в таблице (дефис) изменено: '{original_text}' -> '{new_text}'")

    logger.info(f"task19 выполнена: {history}")
    return doc, history

def task23(doc: Document) -> Tuple[Document, List[str]]:
    """
    Заменяет слово 'Выявлено' на 'Установлено'.
    Возвращает обновленный документ и историю изменений.
    """
    history = ["23. Вместо слова 'Выявлено' использовать 'Установлено'"]
    search_word = 'Выявлено'
    replace_word = 'Установлено'

    for para in doc.paragraphs:
        for run in para.runs:
            if search_word in run.text:
                original_text = run.text
                run.text = run.text.replace(search_word, replace_word)
                run.font.highlight_color = WD_COLOR_INDEX.BLUE
                logger.debug(f"task23 изменено: '{original_text}' -> '{run.text}'")

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if search_word in run.text:
                            original_text = run.text
                            run.text = run.text.replace(search_word, replace_word)
                            run.font.highlight_color = WD_COLOR_INDEX.BLUE
                            logger.debug(f"task23 в таблице изменено: '{original_text}' -> '{run.text}'")
    
    logger.info(f"task23 выполнена: {history}")
    return doc, history

def task24(doc: Document) -> Tuple[Document, List[str]]:
    """
    Заменяет слово 'Сотрудник' на 'Работник'.
    Возвращает обновленный документ и историю изменений.
    """
    history = ["24. Вместо слова 'Сотрудник' использовать 'Работник'"]
    search_word = 'Сотрудник'
    replace_word = 'Работник'

    for para in doc.paragraphs:
        for run in para.runs:
            if search_word in run.text:
                original_text = run.text
                run.text = run.text.replace(search_word, replace_word)
                run.font.highlight_color = WD_COLOR_INDEX.BLUE
                logger.debug(f"task24 изменено: '{original_text}' -> '{run.text}'")

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if search_word in run.text:
                            original_text = run.text
                            run.text = run.text.replace(search_word, replace_word)
                            run.font.highlight_color = WD_COLOR_INDEX.BLUE
                            logger.debug(f"task24 в таблице изменено: '{original_text}' -> '{run.text}'")
    
    logger.info(f"task24 выполнена: {history}")
    return doc, history

def task25(doc: Document) -> Tuple[Document, List[str]]:
    """
    Удаляет множественные пробелы в документе.
    Возвращает обновленный документ и историю изменений.
    """
    history = ["25. Удаление множественных пробелов"]
    pattern = r' {2,}'
    replace_with = ' '

    for para in doc.paragraphs:
        for run in para.runs:
            if re.search(pattern, run.text):
                original_text = run.text
                new_text = re.sub(pattern, replace_with, run.text)
                if original_text != new_text:
                    run.text = new_text
                    run.font.highlight_color = WD_COLOR_INDEX.BLUE
                    logger.debug(f"task25 изменено: '{original_text}' -> '{new_text}'")

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if re.search(pattern, run.text):
                            original_text = run.text
                            new_text = re.sub(pattern, replace_with, run.text)
                            if original_text != new_text:
                                run.text = new_text
                                run.font.highlight_color = WD_COLOR_INDEX.BLUE
                                logger.debug(f"task25 в таблице изменено: '{original_text}' -> '{new_text}'")
    
    logger.info(f"task25 выполнена: {history}")
    return doc, history
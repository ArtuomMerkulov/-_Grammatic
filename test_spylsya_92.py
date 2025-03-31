import re
import pandas as pd
import logging
import socket
from datetime import datetime
import streamlit as st
import docx
from io import BytesIO
from copy import deepcopy
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph
from spylls.hunspell import Dictionary
# Импорт необходимых функций из app_function.py
from app_function import apply_selected_formats, task2func  # Убедитесь, что файл app_function.py доступен
# Импорт необходимых функций из utiles.py
from utiles import convert_docx_to_html  # Предполагается, что файл utiles.py доступен
import os

st.set_page_config(
    page_title="app_EDA_grammer",
    layout="wide",
    initial_sidebar_state="expanded"
)

script_dir = os.path.dirname(os.path.abspath(__file__))
log_file_path = os.path.join(script_dir, 'log.txt')

# Настройка логирования
logging.basicConfig(
    filename=log_file_path,
    level=logging.INFO,
    format='%(asctime)s - %(message)s'
)

logger = logging.getLogger(__name__)

def get_client_info():
    # Получаем имя хоста
    hostname = socket.gethostname()
    # Получаем IP-адрес
    ip_address = socket.gethostbyname(hostname)
    # Получаем порт (в данном случае используем фиктивный порт, так как это не сетевое приложение)
    port = 8514  # Пример порта
    return ip_address, port

allowUsers = ['vw1', 'vw2', 'js3']

auth = Authenticale(
    st.secrets['ldap'],
    st.secrets['session_state_names'],
    st.secrets['auth_cookie']
)

user = auth.login()

if user is None:
    st.caption(
        r':information_source: Username вводится в формате ""',
        unsafe_allow_html=True
    )

if user is not None:
    if user["sAMAccountName"] not in set(allowUsers):
        st.write('''### Ваша учетная запись не найдена.''')
        st.stop()

    auth.createLogoutForm({
        'align': "right",
        'maxWidth': 100,
    })

# Загрузка словаря с кэшированием
@st.cache_resource
def load_dictionary():
    try:
        dictionary = Dictionary.from_files('ru_RU')  # Убедитесь, что файлы словаря находятся в правильном месте
        return dictionary
    except Exception as e:
        st.error(f"Ошибка загрузки словаря: {e}")
        return None

dictionary = load_dictionary()

def match_case(original, corrected):
    """Сохраняет регистр исправленного слова в соответствии с оригинальным словом."""
    if original.isupper():
        return corrected.upper()
    elif original[0].isupper():
        return corrected.capitalize()
    else:
        return corrected

def find_errors_in_text(text, dictionary):
    """Находит орфографические ошибки в тексте, игнорируя числа и аббревиатуры."""
    # Разбиваем текст на слова, игнорируя знаки препинания и одиночные буквы
    tokens = re.findall(r'\b\w{2,}\b', text, re.UNICODE)  # Минимальная длина слова - 2 символа
    errors = []

    for idx, token in enumerate(tokens):
        # Проверяем, что это не аббревиатура (всё заглавное)
        if re.match(r'^[A-ZА-ЯЁ]{2,}$', token): 
            continue
        # Проверяем, есть ли слово в словаре
        if not dictionary.lookup(token):
            suggestions = list(dictionary.suggest(token))
            if suggestions:
                suggestions = [match_case(token, sug) for sug in suggestions]
            else:
                suggestions = []
            errors.append({
                'index': idx,
                'original': token,
                'suggestions': suggestions
            })
    return tokens, errors

def should_add_space(prev_token, current_token):
    """Определяет, нужно ли добавлять пробел перед текущим токеном."""
    if re.match(r'[.,!?;:]', current_token):
        return False
    if re.match(r'[$$$$\{]', prev_token):
        return False
    if re.match(r'[$$$$\}]', current_token):
        return False
    return True

def read_docx(file) -> docx.Document:
    """Читает содержимое .docx файла и возвращает объект Document."""
    try:
        doc = docx.Document(file)
        return doc
    except Exception as e:
        st.error(f"Ошибка чтения .docx файла: {e}")
        return None

def write_docx(doc: docx.Document) -> BytesIO:
    """Сохраняет объект Document в формате .docx и возвращает файл как BytesIO."""
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def copy_run_formatting(source_run, target_run):
    """
    Копирует форматирование из source_run в target_run.
    """
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.strike = source_run.font.strike
    target_run.font.size = source_run.font.size
    if source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb
    target_run.font.name = source_run.font.name
    target_run.style = source_run.style

    # Копируем дополнительные атрибуты форматирования
    rPr_elements = source_run._element.xpath('./w:rPr/*')
    for elem in rPr_elements:
        target_run._element.get_or_add_rPr().append(deepcopy(elem))

def highlight_errors_in_doc(doc, corrections):
    """
    Подсвечивает все найденные ошибки в документе, не изменяя порядок текста.
    Ошибочные слова выделяются желтым фоном.
    Поддерживает как параграфы, так и таблицы.
    """
    for correction in corrections:
        block_type = correction.get('block_type')
        block_index = correction.get('block_index')
        original_word = correction['original']
        
        # Исключаем слова на английском языке, числа и смешанные сочетания букв и цифр
        if (re.match(r'^[a-zA-Z]+$', original_word) or 
            re.match(r'^\d+$', original_word) or 
            re.match(r'^[A-Za-z0-9_]+$', original_word)):
            continue

        if block_type == 'paragraph':
            para = doc.paragraphs[block_index]
            new_runs = []  # Список для новых runs
            for run in para.runs:
                run_text = run.text
                parts = run_text.split(original_word)
                
                # Если оригинальное слово найдено в текущем run
                if len(parts) > 1:  # Проверяем, что слово действительно найдено
                    for i, part in enumerate(parts):
                        if part:
                            new_run = para.add_run(part)  # Добавляем обычный текст
                            copy_run_formatting(run, new_run)  # Копируем форматирование
                            new_runs.append(new_run)  # Сохраняем для последующей обработки
                        if i < len(parts) - 1:
                            # Добавляем ошибочное слово с подсветкой
                            highlighted_run = para.add_run(original_word)
                            copy_run_formatting(run, highlighted_run)
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:fill'), "FFFF00")  # Жёлтый цвет
                            highlighted_run._element.get_or_add_rPr().append(shading_elm)
                            new_runs.append(highlighted_run)

                    # Удаляем старый run после добавления новых
                    run.text = ""
        
        elif block_type == 'table':
            table_index = block_index // 10000
            cell_index = block_index % 10000
            row_idx = cell_index // 100
            col_idx = cell_index % 100

            # Проверяем индексы таблицы и ячейки
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    
                    # Обрабатываем все параграфы в ячейке
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run_text = run.text
                            parts = run_text.split(original_word)
                            
                            # Если слово найдено в этом run
                            if len(parts) > 1:
                                for i, part in enumerate(parts):
                                    if part:
                                        new_run = para.add_run(part)
                                        copy_run_formatting(run, new_run)
                                        
                                    if i < len(parts) - 1:
                                        highlighted_run = para.add_run(original_word)
                                        copy_run_formatting(run, highlighted_run)
                                        shading_elm = OxmlElement('w:shd')
                                        shading_elm.set(qn('w:fill'), "FFFF00")  # Жёлтый цвет
                                        highlighted_run._element.get_or_add_rPr().append(shading_elm)
                                
                                # Удаляем старый run
                                run.text = ""

    return doc

def apply_user_selected_corrections_doc(doc, corrections, selected_corrections):
    """
    Применяет выбранные пользователем исправления к документу.
    Оставляет остальные слова без изменений.
    Поддерживает как параграфы, так и таблицы.
    Возвращает обновленный документ и список применённых исправлений.
    """
    applied_corrections = []

    for correction in corrections:
        block_type = correction.get('block_type')
        block_index = correction.get('block_index')
        original = correction['original']
        suggestions = correction['suggestions']
        corrected_word = suggestions[0] if suggestions else original
        checkbox_key = correction['checkbox_key']

        # Проверяем, выбрал ли пользователь это исправление
        if selected_corrections.get(checkbox_key, False):
            if block_type == 'paragraph':
                # Код для параграфов остается без изменений
                para = doc.paragraphs[block_index]
                for run in para.runs:
                    if original in run.text:
                        parts = run.text.split(original)
                        if len(parts) > 1:  # Проверяем, что слово действительно найдено
                            for i, part in enumerate(parts):
                                if part:
                                    new_run = para.add_run(part)
                                    copy_run_formatting(run, new_run)
                                if i < len(parts) - 1:
                                    # Добавляем исправленное слово с подсветкой
                                    highlighted_run = para.add_run(corrected_word)
                                    copy_run_formatting(run, highlighted_run)
                                    shading_elm = OxmlElement('w:shd')
                                    shading_elm.set(qn('w:fill'), "90EE90")  # Светло-зеленый цвет
                                    highlighted_run._element.get_or_add_rPr().append(shading_elm)
                            run.text = ""
                            applied_corrections.append({
                                'Строка': block_index + 1,
                                'original': original,
                                'corrected': corrected_word
                            })
                            break  # Переходим к следующему исправлению
            
            elif block_type == 'table':
                table_index = block_index // 10000
                cell_index = block_index % 10000
                row_idx = cell_index // 100
                col_idx = cell_index % 100

                try:
                    # Проверяем индексы таблицы и ячейки
                    if table_index < len(doc.tables):
                        table = doc.tables[table_index]
                        
                        if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                            cell = table.rows[row_idx].cells[col_idx]
                            
                            correction_applied = False
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if original in run.text:
                                        parts = run.text.split(original)
                                        if len(parts) > 1:  # Проверяем, что слово действительно найдено
                                            for i, part in enumerate(parts):
                                                if part:
                                                    new_run = para.add_run(part)
                                                    copy_run_formatting(run, new_run)
                                                if i < len(parts) - 1:
                                                    highlighted_run = para.add_run(corrected_word)
                                                    copy_run_formatting(run, highlighted_run)
                                                    shading_elm = OxmlElement('w:shd')
                                                    shading_elm.set(qn('w:fill'), "90EE90")  # Светло-зеленый цвет
                                                    highlighted_run._element.get_or_add_rPr().append(shading_elm)
                                            # Удаляем исходный run
                                            run.text = ""
                                            applied_corrections.append({
                                                'Строка': f"Таблица {table_index + 1}, ячейка ({row_idx + 1}, {col_idx + 1})",
                                                'original': original,
                                                'corrected': corrected_word
                                            })
                                            correction_applied = True
                                            break  # Выходим из цикла run
                                if correction_applied:
                                    break  # Выходим из цикла para
                except IndexError:
                    pass  # Обработка некорректных индексов при необходимости
                except Exception as e:
                    pass  # Обработка других исключений при необходимости

    return doc, applied_corrections

def highlight_corrected_words(doc, corrections):
    """
    Подсвечивает все слова, которые были исправлены.
    Исправленные слова выделяются светло-зеленым фоном.
    Поддерживает как параграфы, так и таблицы.
    """
    for correction in corrections:
        block_type = correction.get('block_type', None)
        block_index = correction.get('block_index', None)
        corrected = correction.get('corrected', '')
        original = correction.get('original', '')
        
        # Если у нас есть информация о строке в формате "Таблица X, ячейка (Y, Z)"
        if 'Строка' in correction and isinstance(correction['Строка'], str) and 'Таблица' in correction['Строка']:
            # Парсим информацию о таблице из строки
            match = re.search(r'Таблица (\d+), ячейка $$(\d+), (\d+)$$', correction['Строка'])
            if match:
                table_index = int(match.group(1)) - 1  # -1 потому что индексация с 0
                row_idx = int(match.group(2)) - 1
                col_idx = int(match.group(3)) - 1
                
                try:
                    table = doc.tables[table_index]
                    cell = table.rows[row_idx].cells[col_idx]
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if corrected in run.text:
                                parts = run.text.split(corrected)
                                if len(parts) > 1:  # Проверяем, что слово действительно найдено
                                    for i, part in enumerate(parts):
                                        if part:
                                            new_run = para.add_run(part)
                                            copy_run_formatting(run, new_run)
                                        if i < len(parts) - 1:
                                            highlighted_run = para.add_run(corrected)
                                            copy_run_formatting(run, highlighted_run)
                                            shading_elm = OxmlElement('w:shd')
                                            shading_elm.set(qn('w:fill'), "90EE90")  # Светло-зеленый цвет
                                            highlighted_run._element.get_or_add_rPr().append(shading_elm)
                                    run.text = ""
                except (IndexError, AttributeError):
                    pass  # Обработка некорректных индексов
        
        elif block_type == 'paragraph' and block_index is not None:
            # Обработка параграфов (оставляем без изменений)
            para = doc.paragraphs[block_index]
            for run in para.runs:
                if corrected in run.text:
                    parts = run.text.split(corrected)
                    if len(parts) > 1:  # Проверяем, что слово действительно найдено
                        for i, part in enumerate(parts):
                            if part:
                                new_run = para.add_run(part)
                                copy_run_formatting(run, new_run)
                            if i < len(parts) - 1:
                                highlighted_run = para.add_run(corrected)
                                copy_run_formatting(run, highlighted_run)
                                shading_elm = OxmlElement('w:shd')
                                shading_elm.set(qn('w:fill'), "90EE90")  # Светло-зеленый цвет
                                highlighted_run._element.get_or_add_rPr().append(shading_elm)
                        run.text = ""
                    
        elif block_type == 'table' and block_index is not None:
            # Обработка таблиц (оставляем логику, но добавляем проверку на существование индексов)
            table_index = block_index // 10000
            cell_index = block_index % 10000
            row_idx = cell_index // 100
            col_idx = cell_index % 100
            try:
                if table_index < len(doc.tables):
                    table = doc.tables[table_index]
                    if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if corrected in run.text:
                                    parts = run.text.split(corrected)
                                    if len(parts) > 1:  # Проверяем, что слово действительно найдено
                                        for i, part in enumerate(parts):
                                            if part:
                                                new_run = para.add_run(part)
                                                copy_run_formatting(run, new_run)
                                            if i < len(parts) - 1:
                                                highlighted_run = para.add_run(corrected)
                                                copy_run_formatting(run, highlighted_run)
                                                shading_elm = OxmlElement('w:shd')
                                                shading_elm.set(qn('w:fill'), "90EE90")  # Светло-зеленый цвет
                                                highlighted_run._element.get_or_add_rPr().append(shading_elm)
                                        run.text = ""
            except IndexError:
                pass  # Обработка некорректных индексов при необходимости
    return doc

def iter_block_items(parent):
    """
    Генератор, перебирающий все параграфы и таблицы в документе, включая вложенные в секции.
    Возвращает элементы в порядке их появления в документе.
    """
    from docx.document import Document

    # Обрабатываем основной контент документа в порядке следования в документе
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)

    # Обрабатываем содержимое секций (например, верхние/нижние колонтитулы)
    if isinstance(parent, Document):
        for sect in parent.sections: 
            for header in sect.header.paragraphs:
                yield header
            for footer in sect.footer.paragraphs:
                yield footer

def convert_docx_to_html(doc, highlight_errors=True, highlight_corrections=False):
    """
    Конвертирует документ docx в HTML с подсветкой ошибок и исправлений.
    
    :param doc: объект Document из python-docx
    :param highlight_errors: если True, подсвечиваются ошибочные слова
    :param highlight_corrections: если True, подсвечиваются исправленные слова
    :return: строка с HTML содержимым
    """
    html = ''
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            paragraph_html = ''
            for run in block.runs:
                run_text = run.text.replace('\n', '<br>')
                # Проверка на подсветку ошибок
                shading = run._element.xpath('.//w:shd')
                if shading:
                    fill = shading[0].get(qn('w:fill'))
                    if fill == "FFFF00" and highlight_errors:
                        # Жёлтая подсветка для ошибок
                        paragraph_html += f'<span style="background-color: #FFFF00">{run_text}</span>'
                    elif fill == "90EE90" and highlight_corrections:
                        # Светло-зелёная подсветка для исправлений
                        paragraph_html += f'<span style="background-color: #90EE90">{run_text}</span>'
                    else:
                        paragraph_html += run_text
                else:
                    paragraph_html += run_text
            html += f'<p>{paragraph_html}</p>'
        elif isinstance(block, Table):
            table_html = '<table style="border-collapse: collapse; border: 1px solid #000;">'
            for row in block.rows:
                table_html += '<tr>'
                for cell in row.cells:
                    cell_html = ''
                    for para in cell.paragraphs:
                        para_html = ''
                        for run in para.runs:
                            run_text = run.text.replace('\n', '<br>')
                            shading = run._element.xpath('.//w:shd')
                            if shading:
                                fill = shading[0].get(qn('w:fill'))
                                if fill == "FFFF00" and highlight_errors:
                                    para_html += f'<span style="background-color: #FFFF00">{run_text}</span>'
                                elif fill == "90EE90" and highlight_corrections:
                                    para_html += f'<span style="background-color: #90EE90">{run_text}</span>'
                                else:
                                    para_html += run_text
                            else:
                                para_html += run_text
                        cell_html += f'<p>{para_html}</p>'
                    table_html += f'<td style="border: 1px solid #000; padding: 5px;">{cell_html}</td>'
                table_html += '</tr>'
            table_html += '</table>'
            html += table_html
    return html

def display_document_with_tables(doc, title, highlight_errors=True, highlight_corrections=False):
    """
    Отображает содержимое документа (абзацы и таблицы) в Streamlit с подсветкой ошибок и исправлений.
    
    :param doc: объект Document из python-docx
    :param title: заголовок раздела
    :param highlight_errors: если True, подсвечиваются ошибочные слова
    :param highlight_corrections: если True, подсвечиваются исправленные слова
    """
    st.subheader(title)
    html = convert_docx_to_html(doc, highlight_errors=highlight_errors, highlight_corrections=highlight_corrections)
    st.markdown(html, unsafe_allow_html=True)

def main():
    ip_address, port = get_client_info()
    login_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    logger.info(f"Вход в систему. IP: {ip_address}, Порт: {port}, Время: {login_time}")
    
    st.title("Автоматический корректор орфографии и форматирования для файлов (.docx)")
    st.write("""
        Этот инструмент позволяет:
        1. Проверить орфографические ошибки в загруженном `.docx` файле.
        2. Выбрать и применить исправления.
        3. Применить форматирующие задачи к исправленному документу.
    """)

    # Компонент загрузки файла
    uploaded_file = st.file_uploader("Выберите файл", type=["docx"])

    # Инициализация флагов состояния
    if 'errors_found' not in st.session_state:
        st.session_state.errors_found = False

    if 'corrections_applied' not in st.session_state:
        st.session_state.corrections_applied = False

    if 'corrections_auto' not in st.session_state:
        st.session_state.corrections_auto = []

    if 'auto_corrected_doc' not in st.session_state:
        st.session_state.auto_corrected_doc = None

    if 'selected_corrections_docx' not in st.session_state:
        st.session_state.selected_corrections_docx = {}

    if 'errors' not in st.session_state:
        st.session_state.errors = []

    if 'last_uploaded_file' not in st.session_state:
        st.session_state.last_uploaded_file = ''

    # 1. Загрузка файла
    if uploaded_file is not None:
        file_type = uploaded_file.name.split('.')[-1].lower()

        # При загрузке нового файла сбрасываем состояния
        if st.session_state.last_uploaded_file != uploaded_file.name:
            st.session_state.errors_found = False
            st.session_state.corrections_applied = False
            st.session_state.corrections_auto = []
            st.session_state.auto_corrected_doc = None
            st.session_state.selected_corrections_docx = {}
            st.session_state.errors = []
            st.session_state.last_uploaded_file = uploaded_file.name

        # Чтение файла
        try:
            if file_type == 'docx':
                original_doc = read_docx(uploaded_file)
            else:
                st.error("Неподдерживаемый формат файла.")
                original_doc = None
        except Exception as e:
            st.error(f"Ошибка чтения файла: {e}")
            original_doc = None

        if file_type == 'docx' and original_doc:
            # 2. Поиск ошибок
            if not st.session_state.errors_found:
                st.subheader("Поиск ошибок:")
                progress_text = "Поиск ошибок..."
                my_bar = st.progress(0, text=progress_text)

                st.session_state.errors = []
                total_blocks = 0
                # Подсчёт количества блоков (параграфы + таблицы)
                for _ in iter_block_items(original_doc):
                    total_blocks += 1

                processed_blocks = 0
                for block_index, block in enumerate(iter_block_items(original_doc)):
                    if isinstance(block, Paragraph):
                        tokens, errors = find_errors_in_text(block.text, dictionary)
                        if errors:
                            for e in errors:
                                st.session_state.errors.append({
                                    'block_type': 'paragraph',
                                    'block_index': block_index,
                                    'index': e['index'],
                                    'original': e['original'],
                                    'suggestions': e['suggestions'],
                                    'checkbox_key': f"checkbox_para_{block_index}_{e['index']}_{e['original']}"
                                })
                    elif isinstance(block, Table):
                        # Обрабатываем каждую ячейку таблицы
                        for table_idx, row in enumerate(block.rows):
                            for col_idx, cell in enumerate(row.cells):
                                # Проверяем текст в каждом параграфе ячейки
                                for para_idx, para in enumerate(cell.paragraphs):
                                    # Проверка на наличие текста или других символов
                                    if para.text.strip():  # Проверяем любой непустой текст
                                        tokens, errors = find_errors_in_text(para.text, dictionary)
                                        if errors:
                                            for e in errors:
                                                # Кодирование блока таблицы и ячейки
                                                encoded_block_index = (block_index * 10000) + (table_idx * 100 + col_idx)
                                                st.session_state.errors.append({
                                                    'block_type': 'table',
                                                    'block_index': encoded_block_index,
                                                    'row': table_idx,
                                                    'col': col_idx,
                                                    'para_idx': para_idx,  # Добавляем индекс параграфа внутри ячейки
                                                    'index': e['index'],
                                                    'original': e['original'],
                                                    'suggestions': e['suggestions'],
                                                    'checkbox_key': f"checkbox_table_{encoded_block_index}_{e['index']}_{e['original']}"
                                                })
                    processed_blocks += 1
                    progress_percentage = int(processed_blocks / total_blocks * 100)
                    my_bar.progress(processed_blocks / total_blocks, text=f"{progress_text} {progress_percentage}% завершено")

                my_bar.empty()  # Удаление прогресс-бара после завершения
                st.session_state.errors_found = True  # Установка флага

            # 3. Оригинальный текст с подсветкой ошибок (для визуализации в Streamlit)
            if st.session_state.errors_found:
                #st.subheader("Оригинальный документ с подсветкой ошибок:")
                original_doc_highlighted = deepcopy(original_doc)

                # Подсвечиваем ошибочные слова
                original_doc_highlighted = highlight_errors_in_doc(original_doc_highlighted, st.session_state.errors)
                st.session_state.original_doc_highlighted = original_doc_highlighted  # Сохраняем подсвеченный документ

                # Визуализация документа с подсветкой ошибок
                display_document_with_tables(
                    st.session_state.original_doc_highlighted, 
                    "Оригинальный документ с подсветкой ошибок:", 
                    highlight_errors=True, 
                    highlight_corrections=False
                )

            # 4. Внесённые изменения (чекбоксы для исправлений)
            if st.session_state.errors_found and st.session_state.errors:
                st.subheader("Выберите исправления:")
                with st.form("corrections_form"):
                    if 'selected_corrections_docx' not in st.session_state:
                        st.session_state.selected_corrections_docx = {}

                    for idx, correction in enumerate(st.session_state.errors):
                        original = correction.get('original', 'Неизвестно')
                        suggestions = correction.get('suggestions', [])
                        if not suggestions:
                            continue  # Пропустить, если нет предложений

                        corrected_word = suggestions[0]  # Берём первое предложение как исправление
                        block_type = correction.get('block_type')
                        block_index = correction.get('block_index')
                        checkbox_key = correction['checkbox_key']

                        if original.lower() != corrected_word.lower():
                            if block_type == 'paragraph':
                                checkbox_label = f"Параграф {block_index + 1}: **{original}** → **{corrected_word}**"
                            elif block_type == 'table':
                                row = correction.get('row', 0)
                                col = correction.get('col', 0)
                                checkbox_label = f"Таблица {(block_index // 10000) + 1}, ячейка ({row + 1}, {col +1}): **{original}** → **{corrected_word}**"

                            if checkbox_key not in st.session_state.selected_corrections_docx:
                                st.session_state.selected_corrections_docx[checkbox_key] = False

                            st.session_state.selected_corrections_docx[checkbox_key] = st.checkbox(
                                checkbox_label,
                                value=st.session_state.selected_corrections_docx[checkbox_key],
                                key=checkbox_key
                            )

                    submitted = st.form_submit_button("Применить выбранные изменения")

                            # 5. Применение исправлений
                if submitted:
                    with st.spinner("Применение выбранных изменений..."):
                        # Применяем выбранные исправления к копии оригинального документа
                        corrected_doc, applied_corrections = apply_user_selected_corrections_doc(
                            deepcopy(original_doc),
                            st.session_state.errors, 
                            st.session_state.selected_corrections_docx
                        )
                        st.session_state.auto_corrected_doc = corrected_doc
                        st.session_state.corrections_applied = True
                
                    # Устанавливаем флаг для отображения сообщения
                    st.session_state.show_correction_success = True
                
                    # Подсветка исправленных слов
                    if st.session_state.auto_corrected_doc:
                        corrected_doc_highlighted = deepcopy(st.session_state.auto_corrected_doc)
                        corrected_doc_highlighted = highlight_corrected_words(corrected_doc_highlighted, applied_corrections)
                        st.session_state.corrected_doc_highlighted = corrected_doc_highlighted

                    # Отображение исправленного документа
                    st.subheader("Исправленный документ с подсветкой исправленных слов:")
                    if 'corrected_doc_highlighted' in st.session_state:
                        display_document_with_tables(
                            st.session_state.corrected_doc_highlighted, 
                            "Исправленный документ с подсветкой исправленных слов",
                            highlight_errors=False, 
                            highlight_corrections=True
                        )

                # Отображаем сообщение только если флаг активен
                if st.session_state.show_correction_success:
                    st.success("Выбранные изменения применены.")

                # 6. Применение форматирующих задач
                st.markdown("### Применение форматирующих задач")
                st.write("Выберите форматирующие задачи, которые необходимо применить к документу.")

                # Состояние для выбранных задач
                selected_tasks = {task: True for task in task2func.keys()}

                # Выбор форматирующих задач с галочками
                for task in task2func.keys():
                    selected_tasks[task] = st.checkbox(task, value=True)

                if st.button("Применить выбранные форматирующие задачи"):
                    tasks_to_apply = [task for task, selected in selected_tasks.items() if selected]
    
                    if tasks_to_apply:
                        with st.spinner("Применение форматирующих задач..."):
                            if st.session_state.auto_corrected_doc:
                                target_doc = st.session_state.auto_corrected_doc
                            else:
                                target_doc = original_doc

                            # Сохраняем результат в session_state
                            st.session_state.formatted_doc, changes, _, _ = apply_selected_formats(target_doc, tasks_to_apply)
                        
                            # Устанавливаем флаг успешного форматирования
                            st.session_state.show_formatting_success = True

                            if changes:
                                st.subheader("Изменения:")
                                for change in changes:
                                    st.write(f"- {change}")
                    else:
                        st.warning("Пожалуйста, выберите хотя бы одну форматирующую задачу.")

                # Отображаем сообщение о форматировании только если флаг активен
                if st.session_state.show_formatting_success:
                    st.success("Форматирующие задачи успешно применены.")

                # Секция скачивания с обработчиком
                if st.session_state.formatted_doc:
                    st.subheader("Скачать Отформатированный Документ:")
                    formatted_file = write_docx(st.session_state.formatted_doc)
                
                    # Обработчик скачивания с сбросом статусов
                    if st.download_button(
                        label="Скачать отформатированный .docx файл",
                        data=formatted_file,
                        file_name="formatted_document.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ):
                        # Сбрасываем статус-сообщения при скачивании
                        st.session_state.show_correction_success = False
                        st.session_state.show_formatting_success = False
                        st.experimental_rerun()

if __name__ == "__main__":
    main()
